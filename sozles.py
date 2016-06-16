# -*- coding: utf-8 -*-
#  name='',
# version='',
# packages=,
# url='',
# license='',
# author='',
# author_email='',
# description=''

from docx import Document
from docx.shared import Pt
from docx.shared import Cm

import os

document = Document('sozlesme.docx')

style = document.styles['Normal']
font = style.font
font.name = 'Book Antiqua'
font.size = Pt(9)

table=document.tables[0]
for cell in table.row_cells(0):
	for paragraph in cell.paragraphs:
		if '[resim123]' in paragraph.text:
			paragraph.text = ' '
			run = cell.paragraphs[0].add_run()
			inline_shape = run.add_picture('./images/001.bmp', width=Cm(2.0))




for paragraph in document.paragraphs:
    if '[adsoyad]' in paragraph.text:
        print paragraph.text
        paragraph.text = u'MAHSUM ŞAHİN'


for table in document.tables:
	print table

	for cell in table.row_cells(6):
		print "xxxxxxxxxxxxx"

		for paragraph in cell.paragraphs:
			if '[adsoyad]' in paragraph.text:
				paragraph.text = u'MAHSUM ŞAHİN'
				paragraph.style = document.styles['Normal']
			if '[kimlikno]' in paragraph.text:
				paragraph.text = u'28233564895'
				paragraph.style = document.styles['Normal']
			if '[adres]' in paragraph.text:
				paragraph.text = u'BARBAROS HAYRETTİN PAŞA M. 1988 SK. NO:6B-16'
				paragraph.style = document.styles['Normal']
			if '[tarih]' in paragraph.text:
				paragraph.text = u'12.06.2016'
				paragraph.style = document.styles['Normal']




document.save('sozlesme2.docx')

os.startfile("sozlesme2.docx")
